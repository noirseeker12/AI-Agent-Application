# 导入Streamlit库，用于创建Web应用
import os

# 导入docx库，用于处理word文档
import docx
import streamlit as st
# 导入PyPDF2库，用于处理pdf文档
from PyPDF2 import PdfReader
from dotenv import load_dotenv
# 导入对话检索链，用于处理对话和检索
from langchain.chains.conversational_retrieval.base import ConversationalRetrievalChain
# 导入递归字符文本分割器，用于将文档分割成小块
from langchain.text_splitter import RecursiveCharacterTextSplitter
# 导入Huggingface文本嵌入模型，用于生成文本嵌入
from langchain_community.embeddings import HuggingFaceEmbeddings
# 导入FAISS向量存储，用于存储和检索文档嵌入
from langchain_community.vectorstores import FAISS
# 导入Document类，用于封装文档内容和元数据
from langchain_core.documents import Document
# 导入OpenAI聊天模型
from langchain_openai import ChatOpenAI

# 设置界面配置、包括标题、图标和布局
st.set_page_config(page_title="文档问答", page_icon=":robot:", layout="wide")

# 设置页面的CSS样式
st.markdown(
    """<style>
    .chat-message {
        display: flex;
        align-items: flex-start;
        gap: 0.75rem;
        padding: 0.75rem 1rem;
        margin-bottom: 0.75rem;
        background: transparent;
        border-radius: 0.75rem;
    }
    
    .chat-message.user { background: transparent; }
    .chat-message.bot { background: transparent; }
    
    .chat-message .avatar {
        flex: 0 0 48px;
        display: flex;
        align-items: center;
        justify-content: center;
    }
    
    .chat-message .avatar img {
        width: 48px;
        height: 48px;
        border-radius: 50%;
        object-fit: cover;
        box-shadow: 0 4px 14px rgba(0,0,0,0.25);
    }
    
    .chat-message .message {
        max-width: calc(100% - 60px);
        padding: 0.9rem 1.1rem;
        border-radius: 0.9rem;
        color: #e9eef7;
        background: linear-gradient(180deg, #4b5568 0%, #3b4354 100%);
        border: 1px solid rgba(255,255,255,0.08);
        box-shadow: 0 6px 20px rgba(0,0,0,0.25);
        line-height: 1.6;
        font-size: 0.95rem;
        white-space: pre-wrap;
        word-break: break-word;
        animation: messageFade 180ms ease-out;
    }
    
    .chat-message.user .message {
        background: linear-gradient(180deg, #2f3543 0%, #242a36 100%);
        border: 1px solid rgba(255,255,255,0.06);
    }
    
    @keyframes messageFade {
        from { opacity: 0; transform: translateY(4px); }
        to { opacity: 1; transform: translateY(0); }
    }
    
    .stDeployButton {
        visibility: hidden;
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    .block-container {
        padding: 2rem 4rem 2rem 4rem;
    }
    
    .st-emotion-cache-16txtl3 {
        padding: 3rem 1.5rem;
    }
    </style>
    
    #"""
    , unsafe_allow_html=True)

# 定义机器人消息模板
bot_template = """
<div class="chat-message bot">
	<div class="avatar">
		<img src="https://cdn.icon-icons.com/icons2/1371/PNG/512/robot02_90810.png" alt="Bot"/>
	</div>
	<div class="message">{{MSG}}</div>
</div>
"""

# 定义用户消息模板
user_template = """
<div class="chat-message user">
	<div class="avatar">
		<img src="https://www.shareicon.net/data/512x512/2015/09/18/103160_man_512x512.png" alt="User"/>
	</div>
	<div class="message">{{MSG}}</div>
</div>
"""

import warnings

warnings.filterwarnings("ignore", category=DeprecationWarning)

os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = "YOUR LANGSMITH_API_KEY"


# 从PDF中提取文本
def get_pdf_text(pdf_docs):
    # 存储提取的文档
    docs = []
    for document in pdf_docs:
        if document.type == "application/pdf":
            # 读取pdf文件
            pdf_reader = PdfReader(document)
            for idx, page in enumerate(pdf_reader.pages):
                docs.append(
                    Document(
                        # 提取页面文本
                        page_content=page.extract_text(),
                        # 添加元数据
                        metadata={"source": f"{document.name} on page {idx}"}
                    )
                )
        elif document.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # 读取word文档
            doc = docx.Document(document)
            for idx, paragraph in enumerate(doc.paragraphs):
                docs.append(
                    Document(
                        # 提取页面文本
                        page_content=paragraph.text,
                        # 添加元数据
                        metadata={"source": f"{document.name} in paragraph {idx}"}
                    )
                )
        elif document.type == "text/plain":
            # 读取纯文本文件
            text = document.getvalue().decode("utf-8")
            docs.append(Document(page_content=text, metadata={"source": document.name}))
    return docs


# 将文档分割成小块文本
def get_text_chunks(docs):
    # 创建文本分割器
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=100)
    # 分割文档
    docs_chunks = text_splitter.split_documents(docs)
    return docs_chunks


# 创建向量存储
def get_vectorstore(doc_chunks):
    # 创建HuggingFace嵌入模型
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )
    # 创建FAISS向量存储
    vectorStore = FAISS.from_documents(doc_chunks, embeddings)
    return vectorStore


# 创建对话检索链
def get_conversation_chain(vectorStore):
    # 创建OpenAI对话模型
    load_dotenv("../../.env")
    llm = ChatOpenAI(
        base_url="https://api.deepseek.com",
        model="deepseek-chat",
        api_key=os.getenv("OPENAI_API_KEY")
    )
    conversation_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        # 使用向量存储作为检索器
        retriever=vectorStore.as_retriever(),
        # 返回源文档
        return_source_documents=True
    )
    return conversation_chain


# 处理用户输入并生成响应
def handle_userinput_pdf(user_question):
    # 获取聊天历史
    chat_history = st.session_state.chat_history
    # 生成响应
    response = st.session_state.conversation(
        {"question": user_question, "chat_history": chat_history}
    )
    # 添加用户问题到聊天历史
    st.session_state.chat_history.append(("user", user_question))
    # 添加机器人回答到聊天历史
    st.session_state.chat_history.append(("assistant", response["answer"]))

    # 显示用户问题
    st.write(
        user_template.replace("{{MSG}}", user_question),
        # 允许嵌入HTML并渲染
        unsafe_allow_html=True
    )

    # 获取源文档
    sources = response["source_documents"]
    # 获取源文档名称
    source_names = set([i.metadata["source"] for i in sources])
    # 合并源文档名称
    src = "\n\n".join(source_names)
    src = f"\n\n> source : {src}"
    message = st.session_state.chat_history[-1]
    # 显示机器人回答和源文档
    st.write(bot_template.replace("{{MSG}}", message[1] + src), unsafe_allow_html=True)


# 显示聊天的历史记录
def show_history():
    # 获取聊天历史
    chat_history = st.session_state.chat_history
    for i, message in enumerate(chat_history):
        if i % 2 == 0:
            # 显示用户消息
            st.write(
                user_template.replace("{{MSG}}", message[1]),
                unsafe_allow_html=True
            )
        else:
            # 显示机器人消息
            st.write(
                bot_template.replace("{{MSG}}", message[1]),
                unsafe_allow_html=True
            )


# 主函数
def main():
    # 显示页面标题
    st.header("Chat with Documents")
    # 初始化会话状态
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    with st.sidebar:
        # 显示侧边标题
        st.title("文档管理")
        # 文件上传控件
        pdf_docs = st.file_uploader(
            "选择文件",
            # 支持的文件类型
            type=["pdf", "txt", "doc", "docx"],
            # 支持多文件上传
            accept_multiple_files=True
        )
        if st.button(
                "处理文档",
                # 设置最后操作为pdf
                on_click=lambda: setattr(st.session_state, "last_action", "pdf"),
                use_container_width=True
        ):
            if pdf_docs:
                # 显示处理中的旋转器
                with st.spinner("Processing"):
                    # 提取pdf、doc、docx文本
                    # chatgpt.pdf 拆分为3个doc
                    # knowledge.txt 拆分为1个doc
                    # new.docx 拆分为37个doc
                    docs = get_pdf_text(pdf_docs)
                    # 分割文本
                    doc_chunks = get_text_chunks(docs)
                    # 创建向量存储
                    vectorStore = get_vectorstore(doc_chunks)
                    # 创建对话链
                    st.session_state.conversation = get_conversation_chain(vectorStore)
            else:
                # 提示用户上传文件
                st.warning("记得上传文件哦~")

        def clear_history():
            # 清空聊天历史
            st.session_state.chat_history = []

        if st.session_state.chat_history:
            # 清空对话按钮
            st.button("清空对话", on_click=clear_history, use_container_width=True)

    with st.container():
        # 获取用户输入
        user_question = st.chat_input("输入点什么~")

    with st.container(height=400):
        # 显示聊天历史
        show_history()
        if user_question:
            if st.session_state.conversation is not None:
                # 处理用户输入
                handle_userinput_pdf(user_question)
            else:
                # 提示用户上传文件
                st.warning("记得上传文件哦~")


if __name__ == '__main__':
    main()
